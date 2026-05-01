import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import logging
import os
import sys
from io import BytesIO

try:
    from PIL import Image
except ImportError:
    Image = None

logger = logging.getLogger(__name__)

def get_resource_path(relative_path):
    """Επιστρέφει την απόλυτη διαδρομή για το αρχείο (συμβατό με PyInstaller)."""
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

def save_to_word(text, stock_name, file_path, chart_image=None, prices=None, stats=None):
    """Εξάγει το κείμενο της ανάλυσης σε έγγραφο Microsoft Word (.docx) διατηρώντας βασική μορφοποίηση."""
    try:
        # Δημιουργία νέου εγγράφου Word
        doc = docx.Document()
        
        # --- Ρύθμιση περιθωρίων εγγράφου ---
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
        
        # --- Προσθήκη Λογότυπου ---
        icon_path = get_resource_path('icon.ico')
        if os.path.exists(icon_path):
            try:
                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                if Image:
                    # Μετατροπή του .ico σε PNG στη μνήμη (το python-docx δεν υποστηρίζει άμεσα .ico)
                    with Image.open(icon_path) as img:
                        img_byte_arr = BytesIO()
                        img.save(img_byte_arr, format='PNG')
                        img_byte_arr.seek(0)
                        logo_para.add_run().add_picture(img_byte_arr, width=Inches(0.6))
                else:
                    logo_para.add_run().add_picture(icon_path, width=Inches(0.6))
            except Exception as e:
                logger.warning(f"Δεν ήταν δυνατή η προσθήκη του λογότυπου: {e}")
        
        # --- Προσθήκη Τίτλου ---
        title = doc.add_heading(f'Ανάλυση: {stock_name}', level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        if chart_image:
            try:
                img_para = doc.add_paragraph()
                img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                img_para.add_run().add_picture(chart_image, width=Inches(6.0))
            except Exception as e:
                logger.warning(f"Δεν ήταν δυνατή η προσθήκη του γραφήματος: {e}")
                
        # --- Πίνακας Τιμών ---
        if prices:
            try:
                table = doc.add_table(rows=1, cols=2)
                try:
                    table.style = 'Table Grid'
                except Exception:
                    pass
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Πηγή (Source)'
                hdr_cells[1].text = 'Τιμή (Price)'
                hdr_cells[0].paragraphs[0].runs[0].bold = True
                hdr_cells[1].paragraphs[0].runs[0].bold = True
                
                for source, price in prices.items():
                    if price and price.strip() and price not in ["---", "N/A", "Σφάλμα", "Φόρτωση..."]:
                        row_cells = table.add_row().cells
                        row_cells[0].text = source
                        row_cells[1].text = price
                doc.add_paragraph()
            except Exception as e:
                logger.warning(f"Δεν ήταν δυνατή η προσθήκη του πίνακα τιμών: {e}")

        # --- Πίνακας Στατιστικών ---
        if stats:
            try:
                doc.add_heading('Βασικά Στατιστικά & Δείκτες Υγείας', level=3)
                stats_table = doc.add_table(rows=1, cols=2)
                try:
                    stats_table.style = 'Table Grid'
                except Exception:
                    pass
                hdr_cells_stats = stats_table.rows[0].cells
                hdr_cells_stats[0].text = 'Δείκτης (Metric)'
                hdr_cells_stats[1].text = 'Τιμή (Value)'
                hdr_cells_stats[0].paragraphs[0].runs[0].bold = True
                hdr_cells_stats[1].paragraphs[0].runs[0].bold = True
                for metric, value in stats.items():
                    if value and value.strip() and value not in ["---", "N/A"]:
                        row_cells = stats_table.add_row().cells
                        row_cells[0].text = metric
                        row_cells[1].text = value
                doc.add_paragraph()
            except Exception as e:
                logger.warning(f"Δεν ήταν δυνατή η προσθήκη του πίνακα στατιστικών: {e}")

        # --- Ρύθμιση Βασικού Στυλ (Normal) ---
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Διαχωρισμός του κειμένου σε γραμμές για ανάλυση του Markdown
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            
            # Κενές γραμμές
            if not line:
                doc.add_paragraph()
                continue
            
            # Επικεφαλίδες (###, ##, #)
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
                continue
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
                continue
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
                continue
            
            # Λίστες (Bullet points & Numbered)
            is_bullet = False
            if line.startswith('- ') or line.startswith('* '):
                line = line[2:]
                p = doc.add_paragraph(style='List Bullet')
                is_bullet = True
            elif re.match(r'^\d+\.\s', line):
                line = re.sub(r'^\d+\.\s', '', line)
                p = doc.add_paragraph(style='List Number')
                is_bullet = True
            else:
                p = doc.add_paragraph()

            # --- Επεξεργασία Bold (**κείμενο**) και Italic (*κείμενο*) ---
            # Χωρίζουμε το κείμενο με βάση τα **
            parts = re.split(r'(\*\*.*?\*\*)', line)
            
            for part in parts:
                if part.startswith('**') and part.endswith('**') and len(part) > 4:
                    # Είναι Bold
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    # Ψάχνουμε για Italic μέσα στα απλά κομμάτια
                    sub_parts = re.split(r'(\*.*?\*)', part)
                    for sub_part in sub_parts:
                        if sub_part.startswith('*') and sub_part.endswith('*') and len(sub_part) > 2:
                            # Είναι Italic
                            run = p.add_run(sub_part[1:-1])
                            run.italic = True
                        else:
                            # Κανονικό κείμενο
                            p.add_run(sub_part)
                            
        # --- Προσθήκη Υποσέλιδου (Footer) ---
        footer = sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "Δημιουργήθηκε από το AI Stock Analyzer Desktop"
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer_para.runs[0].font.size = Pt(9)
        footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128) # Γκρι χρώμα

        doc.save(file_path)
        return True, None
        
    except Exception as e:
        logger.error(f"Σφάλμα κατά την εξαγωγή Word: {e}", exc_info=True)
        return False, str(e)

def save_to_pdf(text, stock_name, file_path, chart_image=None, prices=None, stats=None):
    """Εξάγει το κείμενο σε PDF δημιουργώντας ένα προσωρινό αρχείο Word (απαιτεί docx2pdf και MS Word)."""
    import tempfile
    import os
    try:
        from docx2pdf import convert
    except ImportError:
        return False, "Η βιβλιοθήκη docx2pdf δεν βρέθηκε. Ανοίξτε τερματικό και γράψτε: pip install docx2pdf"
        
    fd, temp_docx = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    
    try:
        success, error = save_to_word(text, stock_name, temp_docx, chart_image, prices, stats)
        if not success:
            return False, f"Σφάλμα δημιουργίας προσωρινού εγγράφου: {error}"
            
        # Αντιμετώπιση του σφάλματος 'tqdm' της docx2pdf όταν δεν υπάρχει ανοιχτή κονσόλα (GUI mode)
        import sys
        from io import StringIO
        original_stdout = sys.stdout
        original_stderr = sys.stderr
        if sys.stdout is None:
            sys.stdout = StringIO()
        if sys.stderr is None:
            sys.stderr = StringIO()
            
        try:
            convert(temp_docx, file_path)
        finally:
            sys.stdout = original_stdout
            sys.stderr = original_stderr
            
        return True, None
    except Exception as e:
        logger.error(f"Σφάλμα κατά την εξαγωγή PDF: {e}", exc_info=True)
        return False, f"Απαιτείται εγκατεστημένο Microsoft Word. Λεπτομέρειες: {str(e)}"
    finally:
        if os.path.exists(temp_docx):
            try:
                os.remove(temp_docx)
            except Exception:
                pass
